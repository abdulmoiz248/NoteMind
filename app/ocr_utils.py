import os
import cv2
import numpy as np
from docx import Document
from PIL import Image
from tqdm import tqdm
from pdf2image import convert_from_path
import pytesseract
import easyocr

reader = easyocr.Reader(['en'], gpu=False)

def enhanceImage(imgPath):
    img = cv2.imread(imgPath)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    # Adaptive thresholding
    adaptive = cv2.adaptiveThreshold(
        gray,
        255,
        cv2.ADAPTIVE_THRESH_MEAN_C,
        cv2.THRESH_BINARY,
        blockSize=11,
        C=10
    )

    prePath = imgPath.replace(".jpg", "_pre.jpg")
    cv2.imwrite(prePath, adaptive)
    return prePath


def extractTextFromDocument(filePath: str, handwritten=False):
    ext = os.path.splitext(filePath)[1].lower()
    extracted = []

    if ext == ".pdf" and handwritten:
        print("🧠 Using EasyOCR for handwritten PDF with OpenCV preprocessing...")
        pages = convert_from_path(filePath, dpi=300)
        for i, img in enumerate(tqdm(pages, desc="Preprocessing + OCR")):
            rawPath = f"temp_page_{i+1}.jpg"
            img.save(rawPath)
            cleanPath = enhanceImage(rawPath)
            result = reader.readtext(cleanPath, detail=0)
            text = "\n".join(result)
            extracted.append({"page": i + 1, "text": text, "image_path": cleanPath})
            os.remove(rawPath)

    elif ext == ".pdf":
        print("📖 Using pytesseract for typed PDF...")
        pages = convert_from_path(filePath, dpi=300)
        for i, img in enumerate(tqdm(pages, desc="Typed PDF OCR")):
            text = pytesseract.image_to_string(img)
            extracted.append({"page": i + 1, "text": text, "image_path": None})

    elif ext in [".png", ".jpg", ".jpeg"]:
        print("🧠 Using EasyOCR with preprocessing on image...")
        cleanPath = enhanceImage(filePath)
        result = reader.readtext(cleanPath, detail=0)
        text = "\n".join(result)
        extracted.append({"page": 1, "text": text, "image_path": cleanPath})

    elif ext == ".docx":
        print("📄 Reading .docx content...")
        doc = Document(filePath)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        extracted.append({"page": 1, "text": text, "image_path": None})

        imgDir = "temp_imgs"
        os.makedirs(imgDir, exist_ok=True)
        imgCount = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                imgCount += 1
                imgData = rel.target_part.blob
                ext = os.path.splitext(rel.target_ref)[1]
                path = os.path.join(imgDir, f"docx_img_{imgCount}{ext}")
                with open(path, "wb") as f:
                    f.write(imgData)
                cleanPath = enhanceImage(path)
                result = reader.readtext(cleanPath, detail=0)
                text = "\n".join(result)
                extracted.append({"page": 1, "text": text, "image_path": cleanPath})

    elif ext == ".txt":
        with open(filePath, "r", encoding="utf-8") as f:
            text = f.read()
        extracted.append({"page": 1, "text": text, "image_path": None})

    else:
        raise ValueError(f"Unsupported file format: {ext}")

    return extracted
